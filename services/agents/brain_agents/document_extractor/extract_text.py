"""Deterministic document content -> text for the document_extractor.

Turns the raw bytes of an uploaded financial document into plain text the
LLM extraction step can read. CSV and plain text decode directly (stdlib);
XLSX spreadsheets are read via openpyxl; DOCX documents are read via
python-docx; PDFs with a text layer are read via pypdf. Scanned/image-only
PDFs and images are a deferred follow-up (OCR needs a model step, which
would break this step's determinism).

This step is deterministic on purpose: the only non-deterministic judgment in
the RFC 0004 pipeline is the LLM field extraction that runs on this text.

RFC F3: this module is pure and synchronous CPU work (pypdf page-by-page,
openpyxl/python-docx zip decompression). The caller (routes.py) is
responsible for running it off the event loop and under a wall-clock
timeout. The bounds here (page count, decompressed size) are the
input-shape guards that must run BEFORE the expensive parse, not after --
unbounded input is otherwise attacker-controlled (raw/server.ts allows up
to 50 MiB uploads).
"""

import io
import zipfile
from typing import Final

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from pypdf.errors import PyPdfError

_CSV_MIMES: Final = frozenset({"text/csv", "application/csv"})
_XLSX_MIME: Final = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_DOCX_MIME: Final = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PDF_MIME: Final = "application/pdf"

# Matches raw/server.ts's 50 MiB artifact cap: reject before parsing rather
# than discover the size mid-parse.
_MAX_INPUT_BYTES: Final = 50 * 1024 * 1024
# Bounds the per-page extract_text() loop; a multi-thousand-page PDF is the
# CPU cost that pins the event loop for minutes (RFC F3).
_MAX_PDF_PAGES: Final = 200
# Zip-bomb guards for XLSX/DOCX: check the sum of each entry's uncompressed
# size (cheap, from the zip central directory) before the OOXML reader
# actually inflates any of it. Both formats are zip containers, so both need
# the same defense.
_MAX_XLSX_UNCOMPRESSED_BYTES: Final = 200 * 1024 * 1024
_MAX_DOCX_UNCOMPRESSED_BYTES: Final = 200 * 1024 * 1024


class UnsupportedDocumentTypeError(Exception):
    """Raised when a document's mime type has no deterministic text extractor."""

    def __init__(self, mime_type: str | None) -> None:
        super().__init__(f"no text extractor for mime type: {mime_type!r}")
        self.mime_type = mime_type


class DocumentTextUnavailableError(Exception):
    """Raised when a supported document type yields no extractable text.

    Covers malformed bytes, password-protected PDFs, and scanned/image-only
    PDFs (no text layer; would need OCR). Raising beats returning "" — empty
    text fed to the LLM extraction step invites hallucinated fields.
    """

    def __init__(self, reason: str) -> None:
        super().__init__(f"no extractable text: {reason}")
        self.reason = reason


class DocumentTooLargeError(Exception):
    """Raised when a document exceeds a size/page bound, before it is parsed."""

    def __init__(self, reason: str) -> None:
        super().__init__(f"document too large: {reason}")
        self.reason = reason


def extract_text(content: bytes, mime_type: str | None) -> str:
    """Return plain text for an uploaded document.

    Raises DocumentTooLargeError when the input exceeds a size/page bound
    (checked before any parsing work runs), UnsupportedDocumentTypeError for
    formats without a deterministic extractor yet (e.g. images), and
    DocumentTextUnavailableError when a supported format carries no readable
    text (scanned/encrypted/broken PDF).
    """
    if len(content) > _MAX_INPUT_BYTES:
        raise DocumentTooLargeError(f"input exceeds {_MAX_INPUT_BYTES // (1024 * 1024)} MB limit")
    mime = (mime_type or "").split(";", 1)[0].strip().lower()
    if mime in _CSV_MIMES:
        return _decode(content)
    if mime == _XLSX_MIME:
        _guard_zip_uncompressed_size(content, _MAX_XLSX_UNCOMPRESSED_BYTES, "XLSX")
        return _xlsx_to_text(content)
    if mime == _DOCX_MIME:
        _guard_zip_uncompressed_size(content, _MAX_DOCX_UNCOMPRESSED_BYTES, "DOCX")
        return _docx_to_text(content)
    if mime == _PDF_MIME:
        return _pdf_to_text(content)
    # text/plain and friends; an empty/missing mime falls back to a utf-8 decode.
    if mime.startswith("text/") or mime == "":
        return _decode(content)
    raise UnsupportedDocumentTypeError(mime_type)


def _decode(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _pdf_to_text(content: bytes) -> str:
    """Read a PDF's text layer, page by page.

    Encrypted PDFs get one empty-password attempt (common for "secured" but
    readable files); anything still locked, malformed, or text-free raises
    DocumentTextUnavailableError so the caller fails loudly instead of
    extracting obligations from a blank page. Page count is checked before
    the per-page extract_text() loop, which is the CPU-heavy part.
    """
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted and not reader.decrypt(""):
            raise DocumentTextUnavailableError("PDF is password-protected")
        page_count = len(reader.pages)
        if page_count > _MAX_PDF_PAGES:
            raise DocumentTooLargeError(
                f"PDF exceeds {_MAX_PDF_PAGES} page limit ({page_count} pages)"
            )
        lines: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text.strip():
                lines.append(f"# page: {number}")
                lines.append(text.strip())
        if not lines:
            raise DocumentTextUnavailableError(
                "PDF has no text layer (scanned/image-only documents need OCR, deferred)"
            )
        return "\n".join(lines)
    except PyPdfError as exc:
        raise DocumentTextUnavailableError(f"unreadable PDF: {exc}") from exc


def _guard_zip_uncompressed_size(content: bytes, max_bytes: int, format_label: str) -> None:
    """Sum each zip entry's declared uncompressed size before the OOXML
    reader (openpyxl/python-docx) inflates any of it. A malformed/non-zip
    file is left for that reader to raise its own, more specific error."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            total_uncompressed = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile:
        return
    if total_uncompressed > max_bytes:
        raise DocumentTooLargeError(
            f"{format_label} uncompressed content exceeds {max_bytes // (1024 * 1024)} MB limit"
        )


def _xlsx_to_text(content: bytes) -> str:
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"# sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                if any(cell.strip() for cell in cells):
                    lines.append("\t".join(cells))
        return "\n".join(lines)
    finally:
        workbook.close()


def _docx_to_text(content: bytes) -> str:
    """Read a Word document's paragraphs and tables, in document order.

    Malformed/non-docx bytes raise DocumentTextUnavailableError rather than
    propagating python-docx's internal exception, matching the PDF path's
    "fail loudly, don't feed garbage to the LLM step" behavior.
    """
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise DocumentTextUnavailableError(f"unreadable DOCX: {exc}") from exc
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_number, table in enumerate(document.tables, start=1):
        rows_with_content = [
            row for row in table.rows if any(cell.text.strip() for cell in row.cells)
        ]
        if not rows_with_content:
            continue
        lines.append(f"# table: {table_number}")
        for row in rows_with_content:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    if not lines:
        raise DocumentTextUnavailableError("DOCX has no extractable text")
    return "\n".join(lines)
